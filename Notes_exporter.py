import os
import re
import pandas as pd
import win32com.client
from datetime import datetime
from docx import Document
from PIL import Image
import pytesseract
import PyPDF2
from openai import OpenAI

# --- Configuration ---
CLIENT_DB_FILE = r'C:\Users\Public\Downloads\Customer info 2.xlsx'
ATTACHMENT_FOLDER = os.path.abspath('attachments')
PATTERN = r"\b\d{6,10}\b"  # Generic regex for invoice, order, PO numbers
DEFAULT_MODEL = os.getenv('OPENAI_MODEL', 'gpt-4.1-nano')

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Ensure attachments folder exists
os.makedirs(ATTACHMENT_FOLDER, exist_ok=True)

# Create and initialize Word document
# file_path: full path where the document will be saved
# account_num and customer_name used for header

def initialize_doc(file_path, account_num, customer_name):
    doc = Document()
    doc.add_heading(f'Account Notes: {customer_name} ({account_num})', level=0)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_paragraph('—' * 40)
    doc.save(file_path)

# Summarization function: 'short' = single email summary; 'long' = conversation summary with numbers and key actions

def summarize(text, summary_type='short'):
    if summary_type == 'short':
        prompt = (
            "Please provide a brief summary (no key actions) of the following email content:\n" + text
        )
    else:
        prompt = (
            "Here is a full email conversation. "
            "For each invoice number, CSO/order number, and PO number mentioned, list the number and give a brief summary of its status. "
            "Then, under 'Key Actions:', list all action items.\n" + text
        )
    response = client.chat.completions.create(
        model=DEFAULT_MODEL,
        messages=[
            {'role': 'system', 'content': 'You are an assistant who summarizes email content.'},
            {'role': 'user', 'content': prompt}
        ],
        max_tokens=800
    )
    return response.choices[0].message.content.strip()

# Extract text from PDF attachments
def extract_pdf(path):
    text = ''
    with open(path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ''
    return text

# Extract text from image attachments via OCR
def extract_image(path):
    return pytesseract.image_to_string(Image.open(path))

# Main routine
def main():
    account_num = input('Enter account number to search: ').strip()
    df = pd.read_excel(CLIENT_DB_FILE, dtype=str)
    # Ensure required columns
    required_cols = ['AccountNumber', 'Email', 'CONTACT_NAME', 'CUSTOMER_NAME']
    for col in required_cols:
        if col not in df.columns:
            print(f"Excel must have column: {col}")
            return
    # Pull associated emails and names
    filter_df = df[df['AccountNumber'] == account_num]
    if filter_df.empty:
        print(f'No entries found for account number {account_num}')
        return
    associated_emails = filter_df['Email'].dropna().unique().tolist()
    # Names for filtering
    associated_names = pd.concat([filter_df['CONTACT_NAME'], filter_df['CUSTOMER_NAME']])
    associated_names = associated_names.dropna().unique().tolist()

    # Get primary customer name for document naming
    customer_name = filter_df['CUSTOMER_NAME'].dropna().iloc[0]
    # Build output path and filename
    output_dir = r'C:\Users\Public\Downloads\Account Email Notes'
    os.makedirs(output_dir, exist_ok=True)
    today_str = datetime.now().strftime('%Y-%m-%d')
    safe_name = re.sub(r"[^\w\- ]", "", customer_name)
    filename = f"{safe_name}_{account_num}_{today_str}.docx"
    file_path = os.path.join(output_dir, filename)

    # Initialize and create Word doc
    initialize_doc(file_path, account_num, customer_name)

    # Setup Outlook
    outlook = win32com.client.Dispatch('Outlook.Application')
    ns = outlook.GetNamespace('MAPI')
    inbox = ns.GetDefaultFolder(6)

    # Build DASL filter
    clauses = [
        f"urn:schemas:httpmail:subject LIKE '%{account_num}%'",
        f"urn:schemas:httpmail:textdescription LIKE '%{account_num}%'"
    ]
    for email in associated_emails:
        clauses += [
            f"urn:schemas:httpmail:senderemailaddress = '{email}'",
            f"urn:schemas:httpmail:to LIKE '%{email}%'",
            f"urn:schemas:httpmail:cc LIKE '%{email}%'"
        ]
    for name in associated_names:
        clauses += [
            f"urn:schemas:httpmail:subject LIKE '%{name}%'",
            f"urn:schemas:httpmail:textdescription LIKE '%{name}%'"
        ]
    filter_query = f'@SQL={' OR '.join(clauses)}'

    try:
        items = inbox.Items.Restrict(filter_query)
    except Exception as e:
        print(f'Error applying filter: {e}')
        items = inbox.Items

    # Collect and summarize
    conversations = {}
    for msg in items:
        try:
            conv_id = getattr(msg, 'ConversationID', None) or msg.ConversationTopic or msg.Subject
            sender = msg.SenderName
            body = msg.Body or ''
            subject = msg.Subject or ''
            sent_on = getattr(msg, 'SentOn', datetime.now())
            sender_addr = getattr(msg, 'SenderEmailAddress', '')
            to = getattr(msg, 'To', '')
            cc = getattr(msg, 'CC', '')
        except:
            continue

        content = f"Subject: {subject}\nBody:\n{body}\nFrom: {sender_addr}\nTo: {to}\nCC: {cc}"
        # Attachments
        for att in msg.Attachments:
            path = os.path.join(ATTACHMENT_FOLDER, att.FileName)
            os.makedirs(os.path.dirname(path), exist_ok=True)
            try:
                att.SaveAsFile(path)
                if path.lower().endswith('.pdf'):
                    content += '\n' + extract_pdf(path)
                elif any(path.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg']):
                    content += '\n' + extract_image(path)
            except Exception as e:
                print(f"Could not save attachment {att.FileName}: {e}")
            finally:
                if os.path.exists(path): os.remove(path)

        # Filter by number, email, or name
        if not (
            re.search(re.escape(account_num), content, re.IGNORECASE)
            or any(email.lower() in content.lower() for email in associated_emails)
            or any(name.lower() in content.lower() for name in associated_names)
        ):
            continue

        # Short summary per email
        summary = summarize(content, summary_type='short')
        conversations.setdefault(conv_id, []).append({
            'sent_on': sent_on,
            'sender': sender,
            'subject': subject,
            'body': body,
            'summary': summary
        })

    # Write details and summaries to Word
    doc = Document(file_path)
    for conv_id, emails in conversations.items():
        emails.sort(key=lambda x: x['sent_on'])
        doc.add_heading(f'Conversation: {conv_id}', level=1)
        for email in emails:
            date_str = email['sent_on'].strftime('%Y-%m-%d %H:%M')
            doc.add_heading(f"Email from {email['sender']} on {date_str}", level=2)
            doc.add_paragraph(email['summary'])
        # Build full conversation text
        convo_text = '\n\n'.join(
            f"From: {e['sender']}\nSubject: {e['subject']}\nBody:\n{e['body']}" for e in emails
        )
        # Detailed summary per conversation
        doc.add_heading('Conversation Summary & Key Actions', level=2)
        convo_summary = summarize(convo_text, summary_type='long')
        doc.add_paragraph(convo_summary)
        doc.add_page_break()

    # Large account-level summary
    all_text = '\n\n'.join(
        f"From: {e['sender']}\nSubject: {e['subject']}\nBody:\n{e['body']}"
        for msgs in conversations.values() for e in msgs
    )
    doc.add_heading('Account-Level Summary & Key Actions', level=1)
    account_summary = summarize(all_text, summary_type='long')
    doc.add_paragraph(account_summary)

    doc.save(file_path)
    print(f'Saved summaries to {file_path}')

if __name__ == '__main__':
    main()
